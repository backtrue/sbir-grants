import sqlite3
from pathlib import Path
from docx import Document
import re


def setup_proposal_db(db_path: Path):
    """
    初始化用於儲存草稿與章節的資料庫表結構。
    """
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS project_sections (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            section_index INTEGER NOT NULL UNIQUE,
            title TEXT NOT NULL,
            content TEXT,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    conn.commit()
    return conn


async def MCP_save_generated_section(section_index: int, title: str, content: str, db_path: str | None = None) -> str:
    """
    讓 Claude 寫完一節計畫書草稿後，將片段精準儲存至本機資料庫。
    """
    try:
        db_base = Path(db_path).resolve() if db_path else Path(__file__).parent.resolve()
        db_file = db_base / "local_skill.db"
        conn = setup_proposal_db(db_file)

        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO project_sections (section_index, title, content, updated_at)
            VALUES (?, ?, ?, CURRENT_TIMESTAMP)
            ON CONFLICT(section_index)
            DO UPDATE SET content = excluded.content, title = excluded.title, updated_at = CURRENT_TIMESTAMP
        ''', (section_index, title, content))

        conn.commit()
        conn.close()
        return f"✅ 成功儲存草稿：第 {section_index} 節「{title}」。"
    except Exception as e:
        return f"❌ 儲存草稿失敗：{str(e)}"


async def MCP_export_proposal_to_word(output_filename: str = "SBIR_Proposal_Draft.docx", db_path: str | None = None) -> str:
    """
    將所有儲存在資料庫中的計畫書章節，依序抓出並組合匯出成一個 Word (.docx) 檔案。
    """
    try:
        db_base = Path(db_path).resolve() if db_path else Path(__file__).parent.resolve()
        db_file = db_base / "local_skill.db"

        if not db_file.exists():
            return "❌ 目前沒有儲存任何草稿紀錄可供匯出。"

        conn = setup_proposal_db(db_file)
        cursor = conn.cursor()

        cursor.execute('SELECT section_index, title, content FROM project_sections ORDER BY section_index ASC')
        sections = cursor.fetchall()
        conn.close()

        if not sections:
            return "❌ 目前沒有儲存任何草稿紀錄可供匯出。"

        doc = Document()
        doc.add_heading('SBIR 計畫書草稿', 0)

        for idx, title, content in sections:
            doc.add_heading(f"第 {idx} 節：{title}", level=1)

            # Simple markdown parsing for the generated text
            # Splitting by double newlines for paragraphs
            paragraphs = str(content).split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                p = doc.add_paragraph()

                # Check for basic Markdown headers (e.g. ## Subtitle)
                header_match = re.match(r'^(#{2,6})\s+(.+)$', para_text)
                if header_match:
                    level = len(header_match.group(1))
                    run = p.add_run(header_match.group(2))
                    run.bold = True
                    try:
                        p.style = doc.styles[f'Heading {min(level, 9)}']
                    except KeyError:
                        try:
                            p.style = doc.styles['Heading 2']
                        except KeyError:
                            pass
                    continue

                # Very naive bolding support `**bold text**`
                # Split by ** and toggle bold
                parts = re.split(r'(\*\*.*?\*\*)', para_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

        current_working_dir = Path.cwd()
        export_path = current_working_dir / output_filename

        # Overwrite if exists
        doc.save(str(export_path))

        return f"✅ 成功組合 {len(sections)} 節草稿，已匯出 Word 檔至：{export_path}"

    except Exception as e:
        return f"❌ 匯出 Word 檔失敗：{str(e)}"


async def MCP_get_all_saved_sections(db_path: str | None = None) -> str:
    """
    讀取所有已儲存的計畫書段落並合併為一個大型字串，提供給 Claude 進行摘要或 Pitch Deck 簡報生成。
    """
    try:
        db_base = Path(db_path).resolve() if db_path else Path(__file__).parent.resolve()
        db_file = db_base / "local_skill.db"

        if not db_file.exists():
            return "❌ 目前沒有儲存任何草稿紀錄。"

        conn = setup_proposal_db(db_file)
        cursor = conn.cursor()

        cursor.execute('SELECT section_index, title, content FROM project_sections ORDER BY section_index ASC')
        sections = cursor.fetchall()
        conn.close()

        if not sections:
            return "❌ 目前沒有儲存任何草稿紀錄。"

        full_document = ""
        for idx, title, content in sections:
            full_document += "\n\n====================\n"
            full_document += f"章節標題：第 {idx} 節：{title}\n"
            full_document += "====================\n\n"
            full_document += str(content)

        return full_document

    except Exception as e:
        return f"❌ 讀取資料失敗：{str(e)}"

if __name__ == "__main__":
    import asyncio
    # Simple manual test
    print(asyncio.run(MCP_save_generated_section(1, "公司簡介", "這是一段測試**粗體**的公司簡介內容。\n\n## 核心理念\n\n創新與實踐！")))
    print(asyncio.run(MCP_save_generated_section(2, "市場分析", "目前市場規模龐大...")))
    print(asyncio.run(MCP_export_proposal_to_word("test_proposal.docx")))
    print(asyncio.run(MCP_get_all_saved_sections())[:100] + "...")
